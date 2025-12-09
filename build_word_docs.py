#!/usr/bin/env python3
"""
Create DOCX with REAL Word tables - Simplified and robust approach
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_styled_table(doc, headers, rows):
    """Add a properly formatted Word table"""
    if not rows:
        rows = [['No data']]
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        # Style header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Blue background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '4472C4')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Add data
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            if col_idx < len(headers):
                cell = table.rows[row_idx].cells[col_idx]
                # Clean markdown
                value = re.sub(r'\*\*(.+?)\*\*', r'\1', value)
                value = re.sub(r'`(.+?)`', r'\1', value)
                cell.text = value
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    return table

def parse_markdown_table(lines):
    """Parse markdown table lines into headers and rows"""
    if len(lines) < 2:
        print(f"        DEBUG: Table too short ({len(lines)} lines)")
        return None, None
    
    # Header
    header_line = lines[0]
    headers = [h.strip() for h in header_line.split('|') if h.strip()]
    
    if not headers:
        print(f"        DEBUG: No headers found in: {header_line[:50]}")
        return None, None
    
    # Data (skip separator line)
    rows = []
    for line in lines[2:]:
        if line.strip().startswith('|') and not line.strip().startswith('|---'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                # Pad row to match header count
                while len(cells) < len(headers):
                    cells.append('')
                rows.append(cells[:len(headers)])
    
    if not rows:
        print(f"        DEBUG: No data rows found. Lines: {len(lines)}")
    
    return headers, rows

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown to DOCX with real tables"""
    doc = Document()
    
    # Configure default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    table_buffer = []
    in_code_block = False
    tables_created = 0
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)
            i += 1
            continue
        
        # Tables
        is_table_line = line.strip().startswith('|') and line.count('|') >= 2
        
        if is_table_line:
            if not table_buffer:
                print(f"        DEBUG: Table start at line {i+1}: {line[:60]}")
            table_buffer.append(line)
            i += 1
            continue
        else:
            # Process accumulated table
            if table_buffer:
                headers, rows = parse_markdown_table(table_buffer)
                print(f"        DEBUG: Parsed - headers={headers is not None}, rows={rows is not None if rows is not None else 'None'}, row_count={len(rows) if rows else 0}")
                if headers and rows is not None:  # Changed: allow empty rows list
                    add_styled_table(doc, headers, rows)
                    doc.add_paragraph()  # Space after table
                    tables_created += 1
                    print(f"      → Table {tables_created}: {len(headers)} columns × {len(rows)} rows")
                table_buffer = []
            
            # Process other content
            if line.startswith('# '):
                doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=3)
            elif line.strip().startswith('- ') or line.strip().startswith('* ') or line.strip().startswith('✅') or line.strip().startswith('❌') or line.strip().startswith('⚙️'):
                text = line.strip().lstrip('-*').strip()
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'`(.+?)`', r'\1', text)
                if text:
                    doc.add_paragraph(text, style='List Bullet')
            elif line.strip() and not line.strip().startswith('---'):
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                text = re.sub(r'`(.+?)`', r'\1', text)
                doc.add_paragraph(text)
            elif not line.strip():
                # Blank line - just skip
                pass
        
        i += 1
    
    # Handle remaining table
    if table_buffer:
        headers, rows = parse_markdown_table(table_buffer)
        if headers and rows:
            add_styled_table(doc, headers, rows)
            tables_created += 1
            print(f"      → Table {tables_created}: {len(headers)} columns × {len(rows)} rows")
    
    doc.save(docx_file)
    return tables_created

def main():
    print("=" * 60)
    print("Creating Professional DOCX Files with Real Word Tables")
    print("=" * 60)
    print()
    
    files = [
        ('DOCUMENTATION_AND_RESOURCES.md', 'DOCUMENTATION_AND_RESOURCES.docx'),
        ('ui_guide.md', 'ui_guide.docx')
    ]
    
    for md_file, docx_file in files:
        print(f"📄 Processing: {md_file}")
        try:
            tables_created = convert_md_to_docx(md_file, docx_file)
            
            # Verify
            doc = Document(docx_file)
            table_count = len(doc.tables)
            print(f"   ✅ Created: {docx_file}")
            print(f"   📊 Tables: {table_count} real Word tables")
            print()
            
        except Exception as e:
            print(f"   ❌ Error: {e}")
            import traceback
            traceback.print_exc()
            print()
    
    print("=" * 60)
    print("✅ DOCX Conversion Complete!")
    print("=" * 60)

if __name__ == '__main__':
    main()
